#!/usr/bin/env python3
# Copyright (c) 2026 정하늘 <ahanaoal@gmail.com> — Licensed under the MIT License.
"""KISA Audit — report.json → report.docx 변환기.

사용:
    python3 tools/render-docx.py <report.json> [-o <out.docx>] [--ip <ip>]
"""

import argparse
import json
import os
import sys
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import remediation guide
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from remediation_guide import REMEDIATION_GUIDE

# ─── 색상 상수 ───
C_HEADER_BG = "2C5F8A"   # 헤더 배경 (진파랑)
C_KEY_BG    = "EBF3FB"   # 키 열 배경 (연파랑)
C_ROW_ALT   = "F5F9FD"   # 짝수행 배경 (매우 연파랑)
C_BORDER    = "5B9BD5"   # 테두리 (파랑)
C_TITLE_FG  = "1E3A5F"   # 항목 제목 글자 (남색)
C_GOOD      = "1A7A3C"   # 양호 (초록)
C_BAD       = "C0392B"   # 취약 (빨강)
C_EV_BEF_BG = "FFF3CD"   # 조치 전 배경 (연황)
C_EV_AFT_BG = "D4EDDA"   # 조치 후 배경 (연초)
C_EV_BEF_BD = "F0C040"   # 조치 전 테두리 (황)
C_EV_AFT_BD = "2E9C5A"   # 조치 후 테두리 (초록)
C_EV_SKIP_BG = "D4EDDA"  # 이미 양호 배경 (연초)
C_EV_SKIP_BD = "2E9C5A"  # 이미 양호 테두리 (초록)
C_EV_MAN_BG  = "FFE0B2"  # 수동 조치 필요 배경 (연주황)
C_EV_MAN_BD  = "E65100"  # 수동 조치 필요 테두리 (주황)


# ─── 유틸 ───

def hex_to_rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def kst_fmt(iso_str):
    if not iso_str:
        return ""
    try:
        dt = datetime.datetime.fromisoformat(iso_str)
        kst = datetime.timezone(datetime.timedelta(hours=9))
        return dt.astimezone(kst).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return iso_str


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=C_BORDER, sz=6):
    """셀 수준 테두리 (style 충돌 없이 확실하게 표시됨)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_all_cell_borders(table, color=C_BORDER, sz=6):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color, sz)


def set_table_width(table, width_cm=16):
    """표 전체 너비 고정."""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int(width_cm * 567)))  # 1 cm ≈ 567 dxa
    tblPr.append(tblW)


def set_col_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcPr.append(tcW)


def cell_run(cell, text, bold=False, size=10, color_hex=None, italic=False, mono=False):
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = hex_to_rgb(color_hex)
    if mono:
        run.font.name = "Courier New"
    return run


# ─── 표 생성 함수 ───

def make_meta_table(doc, rows_data):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Normal Table'
    set_table_width(tbl, 16)
    for i, (k, v) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        set_col_width(c0, 4.5)
        set_col_width(c1, 11.5)
        set_cell_bg(c0, C_KEY_BG)
        set_cell_bg(c1, "FFFFFF")
        set_cell_borders(c0)
        set_cell_borders(c1)
        cell_run(c0, k, bold=True, color_hex=C_TITLE_FG)
        cell_run(c1, v)
    return tbl


def make_summary_table(doc, sum_kisa, mode):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Normal Table'
    set_table_width(tbl, 16)

    col_widths = [3.0, 4.0, 4.0, 5.0]
    headers = ["구분", "전체 항목", "양호", "취약"]

    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        set_col_width(hdr.cells[i], w)
        set_cell_bg(hdr.cells[i], C_HEADER_BG)
        set_cell_borders(hdr.cells[i])
        cell_run(hdr.cells[i], h, bold=True, color_hex="FFFFFF")

    def add_sum_row(label, total, good, bad, bg="FFFFFF"):
        row = tbl.add_row()
        for i, (val, w, clr) in enumerate(zip(
            [label, str(total), str(good), str(bad)],
            col_widths,
            [None, None, C_GOOD, C_BAD]
        )):
            set_col_width(row.cells[i], w)
            set_cell_bg(row.cells[i], C_KEY_BG if i == 0 else bg)
            set_cell_borders(row.cells[i])
            cell_run(row.cells[i], val, bold=(i == 0), color_hex=clr)

    add_sum_row("점검 전",
                sum_kisa.get("total", 0),
                sum_kisa.get("before_good", 0),
                sum_kisa.get("before_bad", 0),
                bg=C_ROW_ALT)
    if mode == "apply":
        add_sum_row("조치 후",
                    sum_kisa.get("total", 0),
                    sum_kisa.get("after_good", 0),
                    sum_kisa.get("after_bad", 0))
    return tbl


def make_item_table(doc, rows_data):
    """항목 기본 정보 표 (2열: 키 | 값)."""
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Normal Table'
    set_table_width(tbl, 16)
    for i, (k, v, bg) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells[0], tbl.rows[i].cells[1]
        set_col_width(c0, 3.5)
        set_col_width(c1, 12.5)
        set_cell_bg(c0, C_KEY_BG)
        set_cell_bg(c1, bg)
        set_cell_borders(c0)
        set_cell_borders(c1)
        cell_run(c0, k, bold=True, color_hex=C_TITLE_FG, size=9)
        cell_run(c1, v, size=9)
    return tbl


def add_evidence_block(doc, label, raw_text, bg_color, border_color):
    """조치 전/후 evidence 블록 — 줄별 paragraph."""
    if not raw_text or not raw_text.strip():
        return

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Normal Table'
    set_table_width(tbl, 16)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_color)
    set_cell_borders(cell, color=border_color, sz=8)
    set_col_width(cell, 16)

    # 라벨 (첫 paragraph)
    lbl_para = cell.paragraphs[0]
    lbl_run = lbl_para.add_run(f"▶ {label}")
    lbl_run.bold = True
    lbl_run.font.size = Pt(9)
    lbl_run.font.color.rgb = hex_to_rgb(border_color)

    # 내용 — 줄별 paragraph
    lines = raw_text.strip().split('\n')
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(8)
        run.font.name = "Courier New"


# ─── 메인 ───

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="report.json 경로")
    ap.add_argument("-o", "--output", help="report.docx 경로 (기본: 같은 디렉터리)")
    ap.add_argument("--ip", help="대상 IP 주소", default="")
    args = ap.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        state = json.load(f)

    mode = state.get("mode", "check")

    doc = Document()

    # 페이지 여백 좁히기
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)

    # 기본 폰트
    doc.styles['Normal'].font.name = '맑은 고딕'
    doc.styles['Normal'].font.size = Pt(10)

    # ─── 타이틀 ───
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("KISA 점검·조치 결과 리포트")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = hex_to_rgb(C_TITLE_FG)

    doc.add_paragraph()

    # ─── 메타데이터 표 ───
    make_meta_table(doc, [
        ("대상 IP 주소", args.ip or "-"),
        ("호스트명",     state.get("host", "")),
        ("운영체제",     state.get("os_pretty", "")),
        ("작업 일자",    kst_fmt(state.get("started_at", ""))),
    ])

    doc.add_page_break()

    # ─── 점검 결과 요약 ───
    doc.add_heading("1. 점검 결과 요약", level=1)
    sum_kisa = state.get("summary_kisa", {})

    # U-01~67 필터링 (요약 + 상세 공용)
    all_items = []
    for it in state.get("items", []):
        code = it.get("code", "")
        if code.startswith("U-"):
            try:
                n = int(code.split("-", 1)[1])
                if 1 <= n <= 67:
                    all_items.append((n, it))
            except (IndexError, ValueError):
                pass
    all_items.sort(key=lambda x: x[0])

    if sum_kisa:
        doc.add_heading("KISA 표준 (U-01 ~ U-67) 통계", level=2)
        make_summary_table(doc, sum_kisa, mode)
        if mode == "apply":
            p = doc.add_paragraph()
            p.add_run(
                f"  자동조치 완료: {sum_kisa.get('applied', 0)}  /  "
                f"이미 양호: {sum_kisa.get('skipped', 0)}  /  "
                f"수동조치 필요: {sum_kisa.get('manual', 0)}  /  "
                f"실패: {sum_kisa.get('failed', 0)}"
            ).font.size = Pt(9)

            # 미조치 항목 목록 (수동 조치 필요 or 취약 유지)
            remaining = [
                it for _, it in all_items
                if it.get('action', '') in ('manual', 'failed', '')
                   and '취약' in str(it.get('after', ''))
            ]
            if remaining:
                doc.add_paragraph()
                hdr_p = doc.add_paragraph()
                hdr_run = hdr_p.add_run("▶ 조치 미완료 항목 (수동 확인 필요)")
                hdr_run.bold = True
                hdr_run.font.size = Pt(10)
                hdr_run.font.color.rgb = hex_to_rgb(C_BAD)
                for it in remaining:
                    bullet_p = doc.add_paragraph(style='List Bullet')
                    bullet_p.paragraph_format.space_before = Pt(0)
                    bullet_p.paragraph_format.space_after  = Pt(2)
                    r = bullet_p.add_run(
                        f"[{it.get('code','')}] {it.get('title','')}  — "
                        f"{it.get('detail','')[:120]}"
                    )
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = hex_to_rgb(C_BAD)
                    r.bold = True

                    code = it.get('code', '')
                    guide = REMEDIATION_GUIDE.get(code)
                    if guide:
                        clean_guide = guide.replace("[조치 방법] ", "").strip()
                        
                        # Add a single-cell table for the command box
                        tbl = doc.add_table(rows=1, cols=1)
                        tbl.style = 'Normal Table'
                        set_table_width(tbl, 15)
                        cell = tbl.rows[0].cells[0]
                        set_cell_bg(cell, C_EV_MAN_BG)
                        set_cell_borders(cell, color=C_EV_MAN_BD, sz=6)
                        set_col_width(cell, 15)
                        
                        # Add guide title
                        p0 = cell.paragraphs[0]
                        p0.paragraph_format.space_before = Pt(2)
                        p0.paragraph_format.space_after = Pt(2)
                        r_lbl = p0.add_run("▶ 추천 조치 방법 및 명령어")
                        r_lbl.bold = True
                        r_lbl.font.size = Pt(8.5)
                        r_lbl.font.color.rgb = hex_to_rgb(C_EV_MAN_BD)
                        
                        # Add guide lines
                        for line in clean_guide.split('\n'):
                            p_line = cell.add_paragraph()
                            p_line.paragraph_format.space_before = Pt(0)
                            p_line.paragraph_format.space_after = Pt(0)
                            r_line = p_line.add_run(line)
                            r_line.font.size = Pt(8)
                            r_line.font.name = "Courier New"
                        
                        # Empty spacer paragraph after table
                        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # ─── 항목별 상세 ───
    doc.add_heading("2. 항목별 상세 결과 (U-01 ~ U-67)", level=1)

    for _, it in all_items:
        code          = it.get("code", "")
        title         = it.get("title", "")
        detail        = it.get("detail", "")
        action_type   = it.get("action", "")   # applied / skipped / manual / failed / ''
        action_method = REMEDIATION_GUIDE.get(code, it.get("action_method", "") or "상세가이드 참고")
        before_val    = it.get("before", "")
        after_val     = it.get("after", "")
        evidence      = it.get("evidence", {}) or {}
        ev_before     = evidence.get("before") or evidence.get("current") or ""
        ev_after      = evidence.get("after") or ""

        # ── 항목 제목 + 상태 뱃지 ──
        h_p = doc.add_paragraph()
        h_p.paragraph_format.space_before = Pt(16)
        h_p.paragraph_format.space_after  = Pt(4)
        h_run = h_p.add_run(f"[{code}]  {title}")
        h_run.bold = True
        h_run.font.size = Pt(12)
        h_run.font.color.rgb = hex_to_rgb(C_TITLE_FG)

        # 상태 뱃지
        badge_map = {
            "applied": ("  ✔ 조치 완료",  C_GOOD),
            "skipped": ("  ✔ 이미 양호",  C_GOOD),
            "manual":  ("  ⚠ 수동조치 필요", C_BAD),
            "failed":  ("  ✖ 실패",        C_BAD),
        }
        if action_type in badge_map:
            badge_txt, badge_color = badge_map[action_type]
            b_run = h_p.add_run(badge_txt)
            b_run.bold = True
            b_run.font.size = Pt(9)
            b_run.font.color.rgb = hex_to_rgb(badge_color)

        # ── 점검 결과 문자열 ──
        if mode == "apply":
            result_str = f"조치 전: {before_val}  →  조치 후: {after_val}"
        else:
            result_str = f"점검 결과: {before_val}"

        # ── 기본 정보 표 ──
        rows_data = [
            ("조치 방법",  action_method, "FFFFFF"),
            ("점검 결과",  result_str,    C_ROW_ALT),
            ("상세 내용",  detail,        "FFFFFF"),
        ]
        make_item_table(doc, rows_data)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ── evidence 블록 (action 유형별 색상) ──
        if action_type == "skipped":
            # 이미 양호 → 초록 단일 블록
            ev_text = ev_before or ev_after
            if ev_text:
                add_evidence_block(doc, "이미 양호 — 조치 불필요 (현재 상태)",
                                   ev_text, C_EV_SKIP_BG, C_EV_SKIP_BD)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif action_type == "applied":
            # 조치 완료 → 노랑(before) + 초록(after)
            if ev_before:
                add_evidence_block(doc, "조치 전 점검 내용",
                                   ev_before, C_EV_BEF_BG, C_EV_BEF_BD)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            if ev_after and ev_before.strip() != ev_after.strip():
                add_evidence_block(doc, "조치 후 점검 내용",
                                   ev_after, C_EV_AFT_BG, C_EV_AFT_BD)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif action_type in ("manual", "failed", ""):
            # 수동 조치 필요 / 실패 → 주황 블록
            ev_text = ev_before or ev_after
            if ev_text:
                add_evidence_block(doc, "수동 조치 필요 — 현재 점검 내용",
                                   ev_text, C_EV_MAN_BG, C_EV_MAN_BD)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)

        else:
            # fallback: check 모드 등
            if ev_before:
                add_evidence_block(doc, "점검 내용", ev_before,
                                   C_EV_BEF_BG, C_EV_BEF_BD)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ─── 저장 ───
    out_docx = args.output or os.path.join(
        os.path.dirname(os.path.abspath(args.input)) or ".", "report.docx"
    )
    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
